# -*- coding: utf-8 -*-
"""
Created on Sun Nov 26 06:37:02 2023
version: 28
@author: ryank, ThongCH, OngWK
"""
import os
import streamlit as st
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from docx import Document
from langchain.text_splitter import CharacterTextSplitter
#from langchain.embeddings import HuggingFaceInstructEmbeddings
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.llms import HuggingFaceHub
from langchain.vectorstores import FAISS
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from htmlTemplates import css
import random
import time
import sklearn
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import CountVectorizer
#import nltk
from nltk.corpus import stopwords
from nltk.corpus import words
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer
from nltk.metrics.distance import edit_distance
#from st_audiorec import st_audiorec
#import transformers
from transformers import BartForConditionalGeneration, BartTokenizer, pipeline
from streamlit_mic_recorder import speech_to_text
import pyttsx3
import urllib.parse
import base64
import docx
from googletrans import Translator
from gtts import gTTS
import tempfile
import subprocess
from autocorrect import Speller
import re




################## CODE for TRANSLATION PURPOSES ###################################
def translateToBMText(text):
    translator = Translator()
    translation = translator.translate(text, dest='ms').text
    return "ms", translation

def translateToMandarin(text):
    translator = Translator()
    translation = translator.translate(text, src='en', dest='zh-cn').text
    return "zh-TW", translation

def translateToEnglish(text):
    translation = "Oops! The content is already in English. But here's the content anyway: " + text
    return "en", translation

def translateText(userQues):
    malay_keywords = ["malay", "translate to malay", "into malay", "to bahasa melayu", "to bahasa malaysia", "melayu"]
    mandarin_keywords = ["mandarin", "translate to mandarin", "into mandarin", "to chinese", "to mandarin", "translate to chinese"]
    english_keywords = ["english","translate to english","into english","to english","english translation","english version","english text"]
    
    userQuesLower = userQues.lower()

    if any(keyword in userQuesLower for keyword in malay_keywords):
        return translateToBMText(st.session_state.fileText)
    elif any(keyword in userQuesLower for keyword in mandarin_keywords):
        return translateToMandarin(st.session_state.fileText)
    elif any(keyword in userQuesLower for keyword in english_keywords):
        return translateToEnglish(st.session_state.fileText)
    else:
        return "en","Sorry, I couldn't determine the target language for translation 😅. I can only translate to Malay or Chinese."

def isTranslationReq(userQues):
    translationKeys = ["translate", "translation", "language", "convert"]
    contextKeys = ["file", "document", "text"]
    questionForms = ["can you", "please", "could you", "would you"]
    userQuesLower = userQues.lower()
    containsTranslationKeys = any(keyword in userQuesLower for keyword in translationKeys)
    containsContextKeys = any(keyword in userQuesLower for keyword in contextKeys)
    containsQuestionForms = any(form in userQuesLower for form in questionForms)
    containsAddRule = "file content translation" in userQuesLower or "to another language" in userQuesLower
    return ((containsTranslationKeys and containsContextKeys) or (containsQuestionForms and containsAddRule))


############################### CODE FOR SUMMARIZATION PURPOSES #########################

def isSummarizationRequest(userQues):
    summarization_keywords = ["summarize", "summarise","summary", "brief", "overview"]
    context_keywords = ["file", "document", "text"]
    question_forms = ["can you", "please", "could you", "would you"]
    user_input_lower = userQues.lower()
    contains_summarizationKey = any(keyword in user_input_lower for keyword in summarization_keywords)
    contains_contextKey = any(keyword in user_input_lower for keyword in context_keywords)
    contains_questionForm = any(form in user_input_lower for form in question_forms)
    contains_addRule = "file summary" in user_input_lower or "provide an overview" in user_input_lower
    return (
        (contains_summarizationKey and contains_contextKey) or
        (contains_questionForm and contains_addRule) or contains_summarizationKey
    )

#NEW CODE
# def summarize_text(text):
#     #text = remove_stopwords(text)
#     model_name = 'facebook/bart-large-cnn'
#     summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
#     result = summarizer(text, max_length=200, min_length=50, do_sample=False)
#     return result[0]['summary_text']

def summarize_text(text, max_chunk_length=700, max_summary_length=130):
    # Initialize the summarization pipeline
    summarizer = pipeline("summarization", model="facebook/bart-large-cnn")

    # Split the text into chunks with a target size
    text_chunks = [text[i:i + max_chunk_length] for i in range(0, len(text), max_chunk_length)]

    # Generate summaries for each chunk
    summaries = []
    for chunk in text_chunks:
        # Use the summarizer pipeline for each chunk
        result = summarizer(chunk, max_length=max_summary_length, min_length=50, do_sample=False)
        summary = result[0]['summary_text']
        #st.write(summary)
        summaries.append(summary)

    # Combine summaries into a final result
    final_summary = " ".join(summaries)
    return final_summary



#############################  CODE FOR DOCUMENT SIMILARITY  ############################
def isDocSimilarityReq(userQues):
    similarityKeys = [
        "document similarity",
        "text similarity",
        "compare documents",
        "analyze document",
        "text comparison",
        "similarity check",
        "content similarity",
        "file resemblance",
        "semantic similarity",
        "matching texts"
    ]
    return any(keyword in userQues.lower() for keyword in similarityKeys)


def chkDocSimilarity(file1, file2):
    docs = [file1, file2]
    countVectorizer = CountVectorizer()
    sparseMatrix = countVectorizer.fit_transform(docs)
    similarityScore = cosine_similarity(sparseMatrix, sparseMatrix)[0, 1]
    return similarityScore










def displayAllChatMsg():
    ind = 0
    for index, message in enumerate(st.session_state.messages):
        if message["role"] == "user":
            avatar = "https://cdn-icons-png.flaticon.com/512/3177/3177440.png"
        else:
            avatar = "https://cdn-icons-png.flaticon.com/512/8943/8943377.png"
        with st.chat_message(message["role"], avatar=avatar):
            st.markdown(message["content"])
            if message["role"] == "assistant":
                st.audio(st.session_state.audioFile[ind], format="audio/mp3")
                ind+=1
            
# def displayAllChatMsg():
#     for index, message in enumerate(st.session_state.messages):
#         if message["role"] == "user":
#             avatar = "https://cdn-icons-png.flaticon.com/512/3177/3177440.png"
#         else:
#             avatar = "https://cdn-icons-png.flaticon.com/512/8943/8943377.png"
        
#         is_last_message = index == len(st.session_state.messages) - 1
        
#         with st.chat_message(message["role"], avatar=avatar):
#             st.markdown(message["content"])
#             if is_last_message:
#                 st.write(is_last_message)
#                 st.audio(st.session_state.audioFile, format="audio/mp3")



def displayChatBotMsg(text, language='en'):
    full_response = ""
    message_placeholder = st.empty()
    for chunk in text.split():
        full_response += chunk + " "
        time.sleep(0.03)
        message_placeholder.markdown(full_response + "▌")
    message_placeholder.markdown(full_response)
    audio_file= text_to_speech(text, language)
    st.session_state.audioFile.append(audio_file)
    st.audio(audio_file, format="audio/mp3")
    

def getPdfText(pdfText):
    text =  ''
    pdf_reader = PdfReader(pdfText)
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def getWordText(docTxt):
    doc = Document(docTxt)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text

def getTextFile(txt):
    text = txt.getvalue().decode('utf-8')
    return text

def getTextChunks(rawText):
    textSplitter = CharacterTextSplitter(
        separator = "\n",
        chunk_size = 950,
        chunk_overlap = 200,
        length_function = len
    )
    txtchunks = textSplitter.split_text(rawText)
    return txtchunks
    
def getVectorStore(textChunks):
    #model_id = 'sentence-transformers/all-MiniLM-L6-v2'
    #modelID = 'sentence-transformers/all-MiniLM-L6-v2'
    #modelID = 'distilbert-base-uncased-distilled-squad'
    #embeddings = HuggingFaceInstructEmbeddings(model_name="hkunlp/instructor-xl")
    #modelKwargs = {'device': 'cpu'}
    embeddings = HuggingFaceEmbeddings(
        model_name='sentence-transformers/all-MiniLM-L6-v2', model_kwargs={'device': 'cpu'}
    )
    vectorStore = FAISS.from_texts(texts=textChunks, embedding = embeddings)
    
    return vectorStore
    

def getConversationChain(vectorStore):
    llm = HuggingFaceHub(repo_id="google/flan-t5-xxl", model_kwargs={"temperature":0.5, "max_length":512})
    memory = ConversationBufferMemory(memory_key='chat_history', return_messages=True)
    conversationChain = ConversationalRetrievalChain.from_llm(
        llm = llm,
        retriever=vectorStore.as_retriever(),
        memory=memory
    )
    return conversationChain

def handleUserInput(userQues):
    ques_ans = [
        {'question': ["Hi","Hi.","Hi!","Helo","Helo.","Helo!","Hi there", "How are you", "Is anyone there?","Hey","Hola", "Hello","Hello!","Hello.", "Good day"],
         'answer' : ["Hello!", "Good to see you again!", "Hi there, how can I help?"] },
        {'question': ["Bye","Bye.","Bye!", "See you later!", "Goodbye!", "Nice chatting to you, bye", "Till next time"], 
         'answer' : ["See you!", "Have a nice day!", "Bye! Come back again soon."] },
        {'question': ["Thanks", "Thank you", "That's helpful", "Awesome, thanks", "Thanks for helping me"], 
         'answer' : ["My pleasure!", "You're Welcome!"] },
        {'question': 
         [
             "What can you do?",
             "What are your capabilities?",
             "Tell me what you can do.",
             "Can you explain your features?",
             "How do you help with files?",
             "What tasks are you good at?",
             "Are there specific file types you handle?",
             "Tell me more about your file-related functions.",
             "What file-related services do you provide?",
             "In what ways can you assist with files?",
             "What kind of tasks can you do?"
         ],
         'answer' : 
             ["Glad you've asked!"
               + " I can answer questions related to the uploaded file,"
               + " translate text to Mandarin and Bahasa Melayu,"
               + " accept both audio and text inputs,"
               + " provide responses in both audio and text,"
               + " accept PDF, DOCX, or TXT files,"
               + " compare document similarity,"
               + " summarize text and"
               + " support general conversation capabilities."
             ]
        }
    ]
    
    found = False  
    largest_index = 0
    largest_cosine_similarity = 0
    largest_jaccard_similarity = 0
    list_ind = 0
    lang = 'en'
    
    for qalist in ques_ans:
        for key, value in qalist.items():
            if key == 'question':
                for item in value:
                    current = item
                    if cosineSimilarity(userQues, current) >= largest_cosine_similarity and jaccardSimilarity(getTokens(userQues), getTokens(current)) >= largest_jaccard_similarity and found == False:
                        found = determineRelationship(userQues, current)
                        largest_index = list_ind
                        largest_cosine_similarity = cosineSimilarity(userQues, current)
                        largest_jaccard_similarity = jaccardSimilarity(getTokens(userQues), getTokens(current))
                        if found:
                            break
        list_ind+=1
    if found == True:
        ans = random.choice(ques_ans[largest_index]['answer'])
    else:
        if isSummarizationRequest(userQues):
            ans = "Certainly! Here's a brief summary: " + summarize_text(st.session_state.fileText)
        elif isTranslationReq(userQues):
            lang, ans = translateText(userQues)
        elif isDocSimilarityReq(userQues):
            ans = "To compare documents, please upload your documents on the left sidebar. Thank you!"
        else:
            try:
                response = st.session_state.conversation({'question': userQues})
                st.session_state.chat_history = response['chat_history']
                #userVoiceText = userVoice[0].upper() + userVoice[1:]
                #new code
                ans = response['answer'][0].upper() + response['answer'][1:]
                really_dk = False
                if ans.lower() == "i don't know," or ans.lower() == "i don't know." or ans.lower() == "i don't know" or ans.lower() == "i do not know" or ans.lower() == "i do not know." or ans.lower() == "i do not know," or ans.lower() == "do not know." or ans.lower() == "do not know," or ans.lower() == "do not know" or  ans.lower() == "no" or ans.lower() == "no." or ans.lower() == "no,"   or ans.lower() == "i don”t know.":
                    really_dk = True
                
                if really_dk:
                    encoded_query = urllib.parse.quote_plus(userQues)
                    google_search_url = f"https://www.google.com/search?q={encoded_query}"
                    ans += " I apologise 😅."
                    ans += f" [Search on Google]({google_search_url})" 
                    response['chat_history'][-1].content = ans
                    response['answer'] = ans
            except:
                    encoded_query = urllib.parse.quote_plus(userQues)
                    google_search_url = f"https://www.google.com/search?q={encoded_query}"
                    ans += "I don't know. I apologise 😅."
                    ans += f" [Search on Google]({google_search_url})" 
        #st.write(isSummarizationRequest(userQues))
    return lang, ans




def cosineSimilarity(userQues, query):
    tokens1 = word_tokenize(userQues)
    tokens2 = word_tokenize(query)
    lemmatizer = WordNetLemmatizer()
    tokens1 = [lemmatizer.lemmatize(token) for token in tokens1]
    tokens2 = [lemmatizer.lemmatize(token) for token in tokens2]

    # Remove stopwords
    stop_words = stopwords.words('english')
    tokens1 = [token for token in tokens1 if token not in stop_words]
    tokens2 = [token for token in tokens2 if token not in stop_words]
    
    tfid_vectorizer = TfidfVectorizer()
    vector1 = tfid_vectorizer.fit_transform(tokens1)
    vector2 = tfid_vectorizer.transform(tokens2)
    similarityBetweenTexts = cosine_similarity(vector1, vector2)
    return similarityBetweenTexts[0,0]

#NEW CODE FOR JACCARD SIMILARITY

def jaccardSimilarity(set_1, set_2):
    intersection = len(set_1.intersection(set_2))
    union = len(set_1.union(set_2))
    return intersection / union if union != 0 else 0

def getTokens(text):
    # Simple tokenization by splitting on spaces
    return set(text.lower().split())

def determineRelationship(text_1, text_2, threshold=0.6):
    tokens_1 = getTokens(text_1)
    tokens_2 = getTokens(text_2)

    similarity = jaccardSimilarity(tokens_1, tokens_2)

    if similarity >= threshold:
        return True
    else:
        return False



def showGeneralResponse(userQues):
    ques_ans = [
        {'question': ["Hi","Hi.","Hi!","Helo","Helo.","Helo!","Hi there", "How are you", "Is anyone there?","Hey","Hola", "Hello","Hello!","Hello.", "Good day"],
         'answer' : ["Hello!", "Good to see you again!", "Hi there, how can I help?"] },
        {'question': ["Bye","Bye.","Bye!", "See you later!", "Goodbye!", "Nice chatting to you, bye", "Till next time"], 
         'answer' : ["See you!", "Have a nice day!", "Bye! Come back again soon."] },
        {'question': ["Thanks", "Thank you", "That's helpful", "Awesome, thanks", "Thanks for helping me"], 
         'answer' : ["My pleasure!", "You're Welcome!"] },
        {'question': 
         [
             "What can you do?",
             "What are your capabilities?",
             "Tell me what you can do.",
             "Can you explain your features?",
             "How do you help with files?",
             "What tasks are you good at?",
             "Are there specific file types you handle?",
             "Tell me more about your file-related functions.",
             "What file-related services do you provide?",
             "In what ways can you assist with files?",
             "What kind of tasks can you do?"
         ],
         'answer' : 
             ["Glad you've asked!"
               + " I can answer questions related to the uploaded file,"
               + " translate text to Mandarin and Bahasa Melayu,"
               + " accept both audio and text inputs,"
               + " provide responses in both audio and text,"
               + " accept PDF, DOCX, or TXT files,"
               + " compare document similarity,"
               + " summarize text and"
               + " support general conversation capabilities."
             ]
        }
    ]
    
    found = False  
    largest_index = 0
    largest_cosine_similarity = 0
    largest_jaccard_similarity = 0
    list_ind = 0
    
    for qalist in ques_ans:
        for key, value in qalist.items():
            if key == 'question':
                for item in value:
                    current = item
                    if cosineSimilarity(userQues, current) >= largest_cosine_similarity and jaccardSimilarity(getTokens(userQues), getTokens(current)) >= largest_jaccard_similarity and found == False:
                        found = determineRelationship(userQues, current)
                        largest_index = list_ind
                        largest_cosine_similarity = cosineSimilarity(userQues, current)
                        largest_jaccard_similarity = jaccardSimilarity(getTokens(userQues), getTokens(current))
                        if found:
                            break
        list_ind+=1
    if found == True:
        response = random.choice(ques_ans[largest_index]['answer'])
    else:
        response = "Sorry, I don't have the answer "
        encoded_query = urllib.parse.quote_plus(userQues)
        google_search_url = f"https://www.google.com/search?q={encoded_query}"
        response += " ("
        response += f" [Search on Google]({google_search_url})" 
        response += ") "
        response += "😅. "
    return response, largest_index                    
                    
#def text_to_speech(text):
#    engine = pyttsx3.init()
#    voices = engine.getProperty('voices')
#    engine.setProperty('voice', voices[1].id)
#    engine.setProperty('voice', 'en-us')
#    engine.setProperty('rate', 180)
#    engine.setProperty('pitch', 0.8)
#    engine.save_to_file(text, "test.mp3")
#    engine.runAndWait()



def text_to_speech(text, language='en'):

    # Specify the language for gTTS
    if language == "en":
        lang = 'en'
    elif language == "ms":
        lang = 'ms'  # Malay
    elif language == "zh-TW":
        lang = 'zh-TW'  # Traditional Chinese
    else:
        lang = 'en'  # Default to English if the language is not recognized

    # Create a gTTS object
    tts = gTTS(text=text, lang=lang, slow=False)

    # Save the audio to a uniquely named file
    timestamp = int(time.time())
    file_path = f"tts_chataudio{timestamp}.mp3"
    tts.save(file_path)
    #st.write("Audio file saved at:", file_path)
    return file_path
          
        


    
    
def get_pdf_display_string(pdf_file):
    base64_pdf = base64.b64encode(pdf_file.getvalue()).decode('utf-8')
    pdf_display_string = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="500" height="1000" type="application/pdf"></iframe>'
    return pdf_display_string

def get_docx_text(uploaded_file):
    doc = docx.Document(uploaded_file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)
        
def simulate_typing(message):
    if message is not None:
        message_placeholder = st.empty()
        full_response = ""
        for char in message:
            full_response += char
            # st.write(char)
            time.sleep(0.01)  
            message_placeholder.markdown(full_response + "▌")

        # Remove the cursor after the message has been fully typed
        message_placeholder.markdown(full_response)

def correctSentence(userQues, exemptWords=['chatbot','chatbot.','chatbot?','chatbot!','chatbots','chatbot,','chatbots,','chatbots?','chatbots.','chatbots!','chatbots,','chatbots/', 'summarise', "hi","hi.","hi!","helo","helo.","helo!","hi there","hey","hola", "hello","hello!","hello.","malay","malay.","malay!","malay?","bahasa","bahasa!","bahasa.","bahasa?","malaysia","malaysia.","malaysia?","malaysia!","melayu","melayu.","melayu!","melayu?"]):
    check = Speller(lang='en')
    def correctWord(word):
        punc = '''!()-[]{};:'"\,<>./?@#$%^&*_~'''
        modified_text = st.session_state.fileText.lower()
        
        for element in modified_text:
            if element in punc:
                modified_text = modified_text.replace(element, "")
        if word.lower() in exemptWords:
            return word
        elif re.sub(r'^\W+|\W+$', '', word).lower() in modified_text.split():
            return word
        else:
            return check(word)
    correctedWords = [correctWord(word) for word in userQues.split()]
    correctedSentence = ' '.join(correctedWords)
    return correctedSentence

            
def main():
    load_dotenv()
    st.set_page_config(page_title="Chat about Your File 📄")
    st.write(css, unsafe_allow_html=True)
    
    page_bg_img = """
    <style>
    [data-testid="stAppViewContainer"] > .main {
        background-image: url("https://static.vecteezy.com/system/resources/previews/021/835/780/original/artificial-intelligence-chatbot-assistance-background-free-vector.jpg");
        background-size: cover;
        background-position: 0px 92px;
        background-repeat: no-repeat;
        background-attachment: scroll;
        background-color: black
        }
    [data-testid='stSidebarUserContent']{
        background-image: url("https://images.unsplash.com/photo-1515549832467-8783363e19b6?q=80&w=1854&auto=format&fit=crop&ixlib=rb-4.0.3&ixid=M3wxMjA3fDB8MHxwaG90by1wYWdlfHx8fGVufDB8fHx8fA%3D%3D");
        background-position: top left;
        background-size: cover;
        }
        
    .stChatFloatingInputContainer {
        background-color: transparent !important;
        }
    [data-testid="stChatMessage"] {
        background-color: #dedede;
        padding: 2%
    }
    
    [data-testid="stChatMessageContent"] p {
        font-size: 1.05em;
        font-family: 'Segoe UI'
        padding: 1em
    }
    
    [data-testid="stChatInput"] {
        border-width: 1px;
        border-radius: 20px;
        border-style: solid;
        border-color: black
    }
    
    
    
    
    
    </style>

    """
    st.markdown("""
        <style>
        .fixed-title {
            position: fixed;
            top: 5%;
            left: 12%;
            width: 100%;
            background-color: #000000;
            z-index: 999;
            padding:  20px 20px;
            text-align: center;
            font-size: 10px;
            line-height: 30px;
        }
        h1 {
            color: white
            }
        .scrollable-content {
            margin-top: 70px;  /* Adjust this margin to fit the height of your title */
        }
        </style>
    """, unsafe_allow_html=True)
    
    st.markdown('<div class="fixed-title"><h1>Chat about Your Uploaded File 📄</h1></div>', unsafe_allow_html=True)
    # Injecting the style into the Streamlit app
    st.markdown(page_bg_img, unsafe_allow_html=True)
    st.write(css, unsafe_allow_html=True)
    #fileText = ""
    
    if "conversation" not in st.session_state:
        st.session_state.conversation = None
        
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = None
    
    if "messages" not in st.session_state:
        st.session_state.messages = []
        
    if 'text_received' not in st.session_state:
        st.session_state.text_received=[]
        
    if 'fileText' not in st.session_state:
        st.session_state.fileText = ""
    
    #NEW CODE
    if 'audioFile' not in st.session_state:
        st.session_state.audioFile = []
    
    language = 'en'
        
    
    #st.title("Chat about Your Uploaded File 📄")
    with st.container():
        displayAllChatMsg()
    #userQues = st.text_input("Ask a question about your documents:")
    
    
    
    #st.write(user_template.replace("{{MSG}}", "hello chatbot"), unsafe_allow_html = True)
    #st.write(bot_template.replace("{{MSG}}", "hello human"), unsafe_allow_html = True)
    
    
    st.sidebar.header('Welcome to FileChatty :wave:')
    
    
    ###############   Upload file  #####################
    st.sidebar.subheader("1. Choose A File for Queries:")
    complete = False
    menu = ["PDF","TXT","DOCX"]
    fileChoice = st.sidebar.selectbox("i) Choose file type",menu)
    
    if fileChoice == "PDF":
        uploadedFile = st.sidebar.file_uploader("ii) Upload your file here", type="pdf", accept_multiple_files=False, key=1)
        
    elif fileChoice == "TXT":
        uploadedFile = st.sidebar.file_uploader("ii) Upload your file here", type="txt", accept_multiple_files=False, key=1)
    else:
        uploadedFile = st.sidebar.file_uploader("ii) Upload your file here", type="docx", accept_multiple_files=False, key=1)
    
    
    
    with st.sidebar:
        if uploadedFile is not None:
        # Display file preview
            st.markdown("<h3 style='color:black;'>File preview</h3>", unsafe_allow_html=True)
            if fileChoice == "PDF":
                st.markdown(get_pdf_display_string(uploadedFile), unsafe_allow_html=True)
            elif fileChoice == "TXT":
                text_content = uploadedFile.read().decode('utf-8')
                st.text_area("Preview", text_content, height=300)
            else:
                document_text = get_docx_text(uploadedFile)
                st.text_area("Preview", document_text, height=300)

        # Get file info
            file_info = f"{uploadedFile.name}_{uploadedFile.size}"

        # Process file only if it's a new file
            if 'last_processed_file' not in st.session_state or st.session_state['last_processed_file'] != file_info:
                with st.spinner("Processing"):
                # Process the file
                    if fileChoice == 'PDF':
                        rawText = getPdfText(uploadedFile)
                    elif fileChoice == 'TXT':
                        rawText = getTextFile(uploadedFile)
                    else:
                        rawText = getWordText(uploadedFile)

                # Store raw text in session state
                    st.session_state.fileText = rawText

                # Process text
                    textChunks = getTextChunks(rawText)
                    vectorStore = getVectorStore(textChunks)
                    st.session_state.conversation = getConversationChain(vectorStore)

            # Update last processed file info
                st.session_state['last_processed_file'] = file_info

            # Indicate completion
                st.write("Upload Complete :white_check_mark:")
                complete = True
                
                    
    st.sidebar.subheader("2. Interact with FileChatty via Text or Voice!")
    st.sidebar.markdown("i) If text, please type your message in the chatbox on the right.")
    st.sidebar.markdown("ii) If voice, please click the button below.")
    
    if st.session_state.messages == []:
        with st.chat_message("assistant", avatar="https://cdn-icons-png.flaticon.com/512/8943/8943377.png"):
            message_placeholder = st.empty()
            greet = "Hi, I'm FileChatty😃. Let's chat via text or voice!"
            full_response = ""
            displayChatBotMsg(greet)
            st.session_state.messages.append({"role": "assistant", "content": greet})
    
    if uploadedFile and complete == True: 
        with st.chat_message("assistant", avatar="https://cdn-icons-png.flaticon.com/512/8943/8943377.png"):
            message_placeholder = st.empty()
            text = "Hi there!\nPlease feel free to ask me anything about " + uploadedFile.name + "!"
            full_response = ""
            displayChatBotMsg(text)
        st.session_state.messages.append({"role": "assistant", "content": text})
        #st.write(bot_template.replace("{{MSG}}", text + uploadedFile.name + "!"), unsafe_allow_html = True)
    
    
    
    
    
    ### INPUT 1: RECEIVE TEXT ###
    
    if userQues := st.chat_input("Chat with FileChatty"):  
        userQues = userQues[0].upper() + userQues[1:]
        userQues = correctSentence(userQues)
        st.session_state.messages.append({"role": "user", "content": userQues})
        with st.chat_message("user", avatar="https://cdn-icons-png.flaticon.com/512/3177/3177440.png"):
            #st.markdown(userQues)
            simulate_typing(userQues)
        
        with st.chat_message("assistant", avatar="https://cdn-icons-png.flaticon.com/512/8943/8943377.png"):
            if uploadedFile:
                language, ans = handleUserInput(userQues)
            else:
                ans, ind = showGeneralResponse(userQues)
                addOn = random.choice(
                     [
                         "😃 Please upload a file to start our conversation. Thank you!",
                         "Please let me have a look at your file! Can't wait to read it 🤩",
                         "🤗 Please upload your file on the left so that I can answer your questions! ",
                     ]
                )
                if ind != 1:
                    ans += ' ' + addOn
            displayChatBotMsg(ans, language)
            st.session_state.messages.append({"role": "assistant", "content": ans})
            
    ### INPUT 2: RECEIVE VOICE
    with st.sidebar:
        userVoice=speech_to_text(start_prompt="Start recording ⏺️",stop_prompt="Stop recording ⏹️", language='en',use_container_width=False,just_once=True,key='STT')
    if userVoice:       
        userVoiceText = userVoice[0].upper() + userVoice[1:]
        st.session_state.text_received.append(userVoiceText)
        st.session_state.messages.append({"role": "user", "content": userVoiceText})
        with st.chat_message("user", avatar="https://cdn-icons-png.flaticon.com/512/3177/3177440.png"):
            st.markdown(userVoiceText)
        
        with st.chat_message("assistant", avatar="https://cdn-icons-png.flaticon.com/512/8943/8943377.png"):
            if uploadedFile:
                language, ans = handleUserInput(userVoiceText)
            else:
                ans, ind = showGeneralResponse(userVoiceText)
                addOn = random.choice(
                     [
                         "😃 Please upload a file to start our conversation. Thank you!",
                         "Please let me have a look at your file! Can't wait to read it 🤩",
                         "🤗 Please upload your file on the left so that I can answer your questions! ",
                     ]
                )
                if ind != 1:
                    ans += ' ' + addOn            
            st.session_state.messages.append({"role": "assistant", "content": ans})
            displayChatBotMsg(ans, language)
            #text_to_speech(ans)   
            
            #st.audio("test.mp3", format="audio/mp3")

    ###### FOR DOCUMENT SIMILARITY ######
    with st.sidebar:
        st.subheader("3. Document Similarity")
        st.markdown("i) Upload 2 documents")
        uploadedFile1 = ''
        uploadedFile2 = ''
        
        menu2 = ["PDF","TXT","DOCX"]
        ## FOR ACCEPTING 1st File ##
        fileChoice1 = st.sidebar.selectbox("Choose file type for File 1",menu2)
        
        if fileChoice1 == "PDF":
            uploadedFile1 = st.file_uploader("Upload your File 1", type="pdf", accept_multiple_files=False, key=2)
            
        elif fileChoice1 == "TXT":
            uploadedFile1 = st.file_uploader("Upload your File 1", type="txt", accept_multiple_files=False, key=2)
        else:
            uploadedFile1 = st.file_uploader("Upload your File 1", type="docx", accept_multiple_files=False, key=2)
        
        
        ## FOR ACCEPTING 2nd File ##        
        fileChoice2 = st.sidebar.selectbox("Choose file type for File 2",menu2)
        
        if fileChoice2 == "PDF":
            uploadedFile2 = st.file_uploader("Upload your File 2", type="pdf", accept_multiple_files=False, key=3)
            
        elif fileChoice2 == "TXT":
            uploadedFile2 = st.file_uploader("Upload your File 2", type="txt", accept_multiple_files=False, key=3)
        else:
            uploadedFile2 = st.file_uploader("Upload your File 2", type="docx", accept_multiple_files=False, key=3)
        
        ## FOR PROCESSING 2 FILES ##
        st.markdown("ii) Click on 'Process Files'")
        
        rawText1 = ''
        rawText2 = ''
        
        if st.button("Process Files", disabled=(uploadedFile1 == None and uploadedFile2 == None)):
            with st.spinner("Processing"):
                #For 1st File
                if fileChoice1 == 'PDF':
                    rawText1 = getPdfText(uploadedFile1)
                elif fileChoice1 == 'TXT':
                    rawText1 = getTextFile(uploadedFile1)
                else:
                    rawText1 = getWordText(uploadedFile1)
                
                #For 2nd File
                if fileChoice2 == 'PDF':
                    rawText2 = getPdfText(uploadedFile2)
                elif fileChoice2 == 'TXT':
                    rawText2 = getTextFile(uploadedFile2)
                else:
                    rawText2 = getWordText(uploadedFile2)
                st.write(f"**Similarity score between the two files:** {chkDocSimilarity(rawText1, rawText2):.2%}")
                pctVal = chkDocSimilarity(rawText1, rawText2)
                if pctVal > 0.8:
                    st.write("**Verdict:** Very High Similarity 😎")
                elif pctVal > 0.6:
                    st.write("**Verdict:** High Similarity 😏")
                elif pctVal > 0.4:
                    st.write("**Verdict:** Moderate Similarity 😐")
                elif pctVal > 0.2:
                    st.write("**Verdict:** Low Similarity 😑")
                else:
                    st.write("**Verdict:** Very Low Similarity 😔")
                
        
    

if __name__ == '__main__':
    main()